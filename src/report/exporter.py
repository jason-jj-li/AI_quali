"""
Word文档导出器
使用python-docx库将报告导出为可编辑的Word文档
"""

import io
from typing import Dict, List, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordExporter:
    """Word文档导出器"""

    def __init__(self):
        """初始化导出器"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

    def export_report(self, report: Dict, output_path: Optional[str] = None) -> bytes:
        """
        导出报告为Word文档

        Args:
            report: 报告数据，包含:
                - title: 标题
                - author: 作者
                - date: 日期
                - abstract: 摘要
                - introduction: 引言
                - methods: 方法
                - results: 结果（主题列表）
                - discussion: 讨论
                - conclusion: 结论
                - references: 参考文献
                - language: 语言
            output_path: 输出文件路径（可选，如果不提供则返回字节）

        Returns:
            Word文档字节或None（如果指定了output_path）
        """
        doc = Document()

        # 设置文档样式
        self._setup_styles(doc, report.get('language', 'zh'))

        # 标题页
        self._add_title_page(doc, report)

        # 摘要
        if report.get('abstract'):
            self._add_abstract(doc, report)

        # 引言
        if report.get('introduction'):
            self._add_section(doc, '1', 'introduction', report['introduction'], report)

        # 方法
        if report.get('methods'):
            self._add_section(doc, '2', 'methods', report['methods'], report)

        # 结果
        if report.get('results'):
            self._add_results(doc, report['results'], report)

        # 讨论
        if report.get('discussion'):
            self._add_section(doc, '4', 'discussion', report['discussion'], report)

        # 结论
        if report.get('conclusion'):
            self._add_section(doc, '5', 'conclusion', report['conclusion'], report)

        # 参考文献
        if report.get('references'):
            self._add_references(doc, report['references'], report)

        # 保存或返回字节
        if output_path:
            doc.save(output_path)
            return None
        else:
            # 返回字节
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

    def _setup_styles(self, doc: Document, language: str):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        if language == 'zh':
            font.name = '宋体'
            font.size = Pt(12)
            # 设置中文字体
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            font.name = 'Times New Roman'
            font.size = Pt(11)

        # 行间距
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)

    def _add_title_page(self, doc: Document, report: Dict):
        """添加标题页"""
        # 标题
        title = report.get('title', '研究报告')
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 作者和日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if report.get('author'):
            p.add_run(f"作者：{report['author']}\n")
        if report.get('date'):
            p.add_run(f"日期：{report['date']}")

        doc.add_page_break()

    def _add_abstract(self, doc: Document, report: Dict):
        """添加摘要"""
        language = report.get('language', 'zh')
        abstract_label = '摘要' if language == 'zh' else 'Abstract'

        doc.add_heading(abstract_label, level=1)

        p = doc.add_paragraph(report['abstract'])
        p.style = 'Normal'

        doc.add_page_break()

    def _add_section(self, doc: Document, number: str, section_id: str,
                    content: str, report: Dict):
        """添加章节"""
        language = report.get('language', 'zh')

        # 章节标题映射
        titles_zh = {
            'introduction': '引言',
            'methods': '方法',
            'discussion': '讨论',
            'conclusion': '结论'
        }

        titles_en = {
            'introduction': 'Introduction',
            'methods': 'Methods',
            'discussion': 'Discussion',
            'conclusion': 'Conclusion'
        }

        titles = titles_zh if language == 'zh' else titles_en
        title = titles.get(section_id, section_id)

        doc.add_heading(f'{number}. {title}', level=1)

        # 添加内容（按段落分割）
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

    def _add_results(self, doc: Document, results: List[Dict], report: Dict):
        """添加结果部分"""
        language = report.get('language', 'zh')
        results_title = '3. 结果' if language == 'zh' else '3. Results'

        doc.add_heading(results_title, level=1)

        for i, theme in enumerate(results, 1):
            # 主题标题
            theme_name = theme.get('theme_name', f'Theme {i}')
            doc.add_heading(theme_name, level=2)

            # 主题描述
            if theme.get('description'):
                doc.add_paragraph(theme['description'])

            # 主题定义
            if theme.get('definition'):
                p = doc.add_paragraph()
                p.add_run('定义：', style='Strong').bold = True
                p.add_run(theme['definition'])

            # 关联编码
            if theme.get('code_names'):
                p = doc.add_paragraph()
                p.add_run('相关编码：', style='Strong').bold = True
                p.add_run(', '.join(theme['code_names']))

            # 典型引用
            if theme.get('quotes'):
                doc.add_heading('典型引用', level=3)

                for quote in theme['quotes']:
                    # 引用文本（斜体）
                    p = doc.add_paragraph()
                    run = p.add_run(f'"{quote["text"]}"')
                    run.italic = True

                    # 引用来源
                    if quote.get('citation'):
                        citation = quote['citation']
                        if language == 'zh':
                            p.add_run(f' —— {citation}')
                        else:
                            p.add_run(f' ({citation})')

    def _add_references(self, doc: Document, references: List[str], report: Dict):
        """添加参考文献"""
        language = report.get('language', 'zh')
        ref_title = '参考文献' if language == 'zh' else 'References'

        doc.add_heading(ref_title, level=1)

        for ref in references:
            doc.add_paragraph(ref, style='List Number')

    def export_to_file(self, report: Dict, filename: str) -> str:
        """
        导出到文件

        Args:
            report: 报告数据
            filename: 文件名

        Returns:
            完整的文件路径
        """
        if not filename.endswith('.docx'):
            filename += '.docx'

        self.export_report(report, filename)
        return filename

    def create_downloadable_filename(self, report_title: str) -> str:
        """
        创建可下载的文件名

        Args:
            report_title: 报告标题

        Returns:
            文件名
        """
        # 清理标题中的非法字符
        safe_title = "".join(c for c in report_title if c.isalnum() or c in (' ', '-', '_'))
        safe_title = safe_title.strip()

        # 添加日期戳
        date_str = datetime.now().strftime('%Y%m%d')
        return f"{safe_title}_{date_str}.docx"


def test_word_export():
    """测试Word导出功能"""
    if not DOCX_AVAILABLE:
        print("python-docx not available")
        return

    # 测试数据
    test_report = {
        'title': '测试研究报告',
        'author': '研究者',
        'date': '2024-01-15',
        'abstract': '这是一个测试摘要...',
        'introduction': '这是引言部分。\n\n引言的第二段。',
        'methods': '这是方法部分。\n\n使用质性研究方法。',
        'results': [
            {
                'theme_name': '主题一',
                'description': '这是第一个主题的描述。',
                'definition': '主题一的学术定义。',
                'code_names': ['编码1', '编码2'],
                'quotes': [
                    {
                        'text': '这是第一个典型引用。',
                        'citation': '[文档1, P01]'
                    },
                    {
                        'text': '这是第二个典型引用。',
                        'citation': '[文档2, P02]'
                    }
                ]
            }
        ],
        'discussion': '这是讨论部分。',
        'conclusion': '这是结论部分。',
        'references': ['参考文献1', '参考文献2'],
        'language': 'zh'
    }

    exporter = WordExporter()
    output_file = exporter.export_to_file(test_report, 'test_report')
    print(f"Report exported to: {output_file}")


if __name__ == '__main__':
    test_word_export()
